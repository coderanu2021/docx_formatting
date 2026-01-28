from docx import Document
from app import process_docx
import os

def create_test_docx(path):
    doc = Document()
    doc.add_paragraph("ARTICLE MAIN TITLE") # Main Heading (1-column)
    doc.add_paragraph("Keywords: Automation, Pharmaceutical") # Keywords (2-column)
    doc.add_paragraph("1. Cost reduction") # Subheading (2-column)
    doc.add_paragraph("Automation can reduce labour cost thus reducing operation cost.") # Para 1
    doc.add_paragraph("INNER HEADING IN UPPERCASE") # Inner heading (should be 2-column)
    doc.add_paragraph("As much of the tasks get automated humans will only be needed.") # Para 2
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header 1"
    table.cell(0, 1).text = "Header 2"
    table.cell(1, 0).text = "Data 1"
    table.cell(1, 1).text = "Data 2"
    
    doc.save(path)

def verify_output(path):
    doc = Document(path)
    print(f"Verifying {path}...")
    
    # Check paragraphs and tables
    found_bold_heading = False
    found_keywords = False
    para_count = 0
    found_table = False
    
    for block in doc.paragraphs:
        if "1. Cost reduction" in block.text:
            # Check if bold
            for run in block.runs:
                if run.bold:
                    found_bold_heading = True
        if "Keywords" in block.text:
            for run in block.runs:
                if run.bold:
                    found_keywords = True
        if "Automation" in block.text or "As much" in block.text:
            para_count += 1
            
    if doc.tables:
        found_table = True
        
    print(f"Found bold heading: {found_bold_heading}")
    print(f"Found keywords: {found_keywords}")
    print(f"Paragraphs found: {para_count}")
    print(f"Table found: {found_table}")
    
    section_count = len(doc.sections)
    print(f"Sectons found: {section_count}")
    
    # section_count should be exactly 2: 
    # 1. First section (1-col) for title
    # 2. Second section (2-col) for the rest
    return found_bold_heading and found_keywords and para_count >= 2 and found_table and section_count == 2

if __name__ == "__main__":
    test_input = "test_input.docx"
    test_output = "test_output.docx"
    
    create_test_docx(test_input)
    if process_docx(test_input, test_output):
        if verify_output(test_output):
            print("Verification SUCCESS")
        else:
            print("Verification FAILED")
    else:
        print("Processing FAILED")
    
    # Cleanup
    if os.path.exists(test_input): os.remove(test_input)
    # Don't remove output yet so user can see it if they want
