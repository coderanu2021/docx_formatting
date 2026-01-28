from flask import Flask, request, render_template, send_file, flash, redirect, url_for
import os
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.table import Table, _Cell
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tempfile
import zipfile
import shutil
import io
import re
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Create upload directory if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

ALLOWED_EXTENSIONS = {'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/favicon.ico')
def favicon():
    return '', 204  # No content response for favicon

def add_page_numbers_and_footer(section, footer_text=""):
    """Add page numbers and custom footer text to the document"""
    try:
        # Create footer
        footer = section.footer
        footer_para = footer.paragraphs[0]
        
        # Clear existing content
        footer_para.clear()
        
        if footer_text.strip():
            # Set paragraph format with tab stops
            pPr = footer_para._element.get_or_add_pPr()
            
            # Add tab stops - one for right alignment
            tabs = OxmlElement('w:tabs')
            tab = OxmlElement('w:tab')
            tab.set(qn('w:val'), 'right')
            tab.set(qn('w:pos'), '9360')  # 6.5 inches in twentieths of a point
            tabs.append(tab)
            pPr.append(tabs)
            
            # Add page number on the left
            page_run = footer_para.add_run("Page ")
            page_run.font.size = Inches(0.14)
            
            # Add page number field
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            page_run._element.append(fldChar1)
            page_run._element.append(instrText)
            page_run._element.append(fldChar2)
            
            # Add tab character to move to right side
            tab_run = footer_para.add_run("\t")
            
            # Add custom footer text on the right
            footer_run = footer_para.add_run(footer_text)
            footer_run.font.size = Inches(0.14)
            
        else:
            # Just page number
            page_num_run = footer_para.add_run("Page ")
            page_num_run.font.size = Inches(0.14)
            
            # Add page number field
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            page_num_run._element.append(fldChar1)
            page_num_run._element.append(instrText)
            page_num_run._element.append(fldChar2)
        
    except Exception as e:
        print(f"Error adding page numbers and footer: {e}")
        pass

def configure_section(section, columns=1, footer_text=""):
    """Configure margins, columns, and footer for a section"""
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    
    # Set up columns
    sectPr = section._sectPr
    # Clear existing columns if any
    cols_elements = sectPr.xpath('./w:cols')
    for cols in cols_elements:
        sectPr.remove(cols)
        
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), str(columns))
    cols.set(qn('w:space'), '720')  # 0.5 inch space between columns
    sectPr.append(cols)
    
    # Add page numbers and footer
    add_page_numbers_and_footer(section, footer_text)

def get_optimal_image_size(image_path, max_width_inches=4.5):
    """Calculate optimal image size based on aspect ratio and available space"""
    try:
        from PIL import Image
        with Image.open(image_path) as img:
            width, height = img.size
            aspect_ratio = height / width
            
            # Calculate optimal width (max 4.5 inches to fit better in document)
            optimal_width = min(max_width_inches, 4.5)
            optimal_height = optimal_width * aspect_ratio
            
            # If height is too large, adjust based on height constraint
            if optimal_height > 3:  # Max 3 inches height
                optimal_height = 3
                optimal_width = optimal_height / aspect_ratio
            
            # Ensure minimum readable size
            if optimal_width < 2:
                optimal_width = 2
                optimal_height = optimal_width * aspect_ratio
            
            return Inches(optimal_width), Inches(optimal_height)
    except Exception as e:
        print(f"PIL not available or error processing image: {e}")
        # Default smaller size if PIL is not available
        return Inches(4), Inches(3)

def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph.
    *parent* would most commonly be a Document object.
    """
    from docx.document import Document as DocumentClass
    if isinstance(parent, DocumentClass):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Unknown parent type")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def copy_table(src_table, dest_doc):
    """Copy a table from source to destination document with preservation of text and basic formatting"""
    new_table = dest_doc.add_table(rows=len(src_table.rows), cols=len(src_table.columns))
    new_table.style = src_table.style
    
    # Copy alignment and column widths if possible
    for r_idx, row in enumerate(src_table.rows):
        for c_idx, cell in enumerate(row.cells):
            new_cell = new_table.rows[r_idx].cells[c_idx]
            # Copy text and formatting by copying paragraphs
            first_p = True
            for p in cell.paragraphs:
                if first_p:
                    new_p = new_cell.paragraphs[0]
                    new_p.text = "" # clear default
                    first_p = False
                else:
                    new_p = new_cell.add_paragraph()
                
                for run in p.runs:
                    new_run = new_p.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    if run.font.size:
                        new_run.font.size = run.font.size
                    if run.font.name:
                        new_run.font.name = run.font.name

def process_docx(input_path, output_path, footer_text=""):
    """Process DOCX file to make headings bold and arrange content in 2 columns"""
    try:
        # Open the document
        doc = Document(input_path)
        
        # Create new document for output
        new_doc = Document()
        
        # Configure initial section as 1-column for title
        configure_section(new_doc.sections[0], columns=1, footer_text=footer_text)
        current_columns = 1
        title_found = False
        
        in_references = False
        
        # Extract images from the original document
        image_dict = extract_images_from_docx(input_path)
        
        # Process all block items (paragraphs and tables) in order
        for block in iter_block_items(doc):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                
                # Check for images in this paragraph
                has_images = False
                for run in block.runs:
                    if 'graphicData' in run._element.xml or 'pic:pic' in run._element.xml:
                        has_images = True
                        break
                
                if has_images:
                    # Images usually follow current layout
                    new_para = new_doc.add_paragraph()
                    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    image_added = False
                    # Try to add images from extracted files
                    if image_dict:
                        try:
                            image_key = list(image_dict.keys())[0]
                            image_path = image_dict[image_key]
                            new_run = new_para.add_run()
                            width, height = get_optimal_image_size(image_path)
                            new_run.add_picture(image_path, width=width, height=height)
                            image_added = True
                            del image_dict[image_key]
                        except Exception as img_error:
                            print(f"Could not add image: {img_error}")
                    
                    if not image_added:
                        new_para.add_run("[Image placeholder]")
                    
                    if text:
                        caption_para = new_doc.add_paragraph(text)
                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in caption_para.runs:
                            run.font.size = Inches(0.12)
                            run.italic = True
                    
                    new_doc.add_paragraph()
                    continue

                if not text:
                    continue

                # Check if this is references section
                if 'REFERENCES' in text.upper() or 'BIBLIOGRAPHY' in text.upper():
                    in_references = True
                    heading_para = new_doc.add_paragraph()
                    heading_run = heading_para.add_run(text)
                    heading_run.bold = True
                    heading_run.font.size = Inches(0.16)
                    heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    new_doc.add_paragraph()
                    continue

                # Handle references
                if in_references and text and not text.isupper():
                    ref_para = new_doc.add_paragraph()
                    ref_para.style = 'List Number'
                    ref_para.add_run(text)
                    continue

                # Improved heading detection
                # Main Heading: UPPERCASE text with more than 1 word
                is_uppercase_heading = text.isupper() and len(text.split()) > 1
                
                # Special Blocks: Abstract or Keywords (Stay in 2 columns but are bold)
                is_special_block = (
                    text.upper().startswith('ABSTRACT') or 
                    text.upper().startswith('KEYWORDS')
                )
                
                # Subheading: Numbered lists or specific keywords
                is_subheading = (
                    re.match(r'^(\d+|[A-ZIVX]+)[\.\)]\s+', text) or
                    (len(text.split()) <= 8 and any(word in text.upper() for word in ['INTRODUCTION', 'CONCLUSION', 'CHAPTER', 'SECTION', 'METHODOLOGY', 'RESULT', 'DISCUSSION', 'ABSTRACT', 'KEYWORDS'])) or
                    text.endswith(':') or
                    (text.isupper() and len(text.split()) <= 1) # Single word uppercase like "AIM"
                )

                if is_uppercase_heading and not title_found: 
                    # ONLY the FIRST uppercase heading spans 1 row (full width)
                    title_found = True
                    # Switch to 1 column if not already
                    if current_columns != 1:
                        new_section = new_doc.add_section(WD_SECTION.CONTINUOUS)
                        configure_section(new_section, columns=1, footer_text=footer_text)
                        current_columns = 1
                    
                    heading_para = new_doc.add_paragraph()
                    heading_run = heading_para.add_run(text)
                    heading_run.bold = True
                    heading_run.font.size = Inches(0.18)
                    heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    new_doc.add_paragraph()
                elif is_uppercase_heading or is_special_block or is_subheading:
                    # Subsequent uppercase headings, Abstract/Keywords, and Subheadings stay in 2 columns
                    if current_columns != 2:
                        new_section = new_doc.add_section(WD_SECTION.CONTINUOUS)
                        configure_section(new_section, columns=2, footer_text=footer_text)
                        current_columns = 2
                        
                    heading_para = new_doc.add_paragraph()
                    heading_run = heading_para.add_run(text)
                    heading_run.bold = True
                    heading_run.font.size = Inches(0.14)
                    
                    if is_special_block:
                        heading_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    else:
                        heading_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                    new_doc.add_paragraph()
                else:
                    # Normal paragraph - ensure in 2 columns
                    if current_columns != 2:
                        new_section = new_doc.add_section(WD_SECTION.CONTINUOUS)
                        configure_section(new_section, columns=2, footer_text=footer_text)
                        current_columns = 2
                        
                    new_para = new_doc.add_paragraph(text)
                    new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    new_para.paragraph_format.space_after = Inches(0.1)
            
            elif isinstance(block, Table):
                # Copy table to new document
                copy_table(block, new_doc)
                new_doc.add_paragraph() # spacing after table

        # Clean up extracted images
        for img_path in image_dict.values():
            try:
                os.remove(img_path)
            except: pass
        
        new_doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error processing document: {str(e)}")
        import traceback
        traceback.print_exc()
        return False
def extract_images_from_docx(docx_path):
    """Extract images from DOCX file and return dict of image paths"""
    image_dict = {}
    temp_dir = tempfile.mkdtemp()
    
    try:
        # Extract DOCX as ZIP
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
        
        # Find media folder
        media_path = os.path.join(temp_dir, 'word', 'media')
        if os.path.exists(media_path):
            for filename in os.listdir(media_path):
                if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                    image_path = os.path.join(media_path, filename)
                    # Copy to temp location
                    temp_image_path = os.path.join(tempfile.gettempdir(), f"temp_{filename}")
                    shutil.copy2(image_path, temp_image_path)
                    image_dict[filename] = temp_image_path
    
    except Exception as e:
        print(f"Error extracting images: {e}")
    
    finally:
        # Clean up temp directory
        try:
            shutil.rmtree(temp_dir)
        except:
            pass
    
    return image_dict


@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        flash('No file selected')
        return redirect(request.url)
    
    file = request.files['file']
    footer_text = request.form.get('footer_text', '')  # Get footer text from form
    
    if file.filename == '':
        flash('No file selected')
        return redirect(request.url)
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)
        
        # Create output filename
        name, ext = os.path.splitext(filename)
        output_filename = f"{name}_processed{ext}"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
        
        # Process the document
        if process_docx(input_path, output_path, footer_text):
            # Clean up input file
            os.remove(input_path)
            return send_file(output_path, as_attachment=True, download_name=output_filename)
        else:
            flash('Error processing document')
            os.remove(input_path)
            return redirect(url_for('index'))
    else:
        flash('Invalid file type. Please upload a .docx file')
        return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True)